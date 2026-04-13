from __future__ import annotations

from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "documentation_1.docx"
BACKUP_PATH = ROOT / "documentation_1_backup.docx"
SCREENSHOTS_DIR = ROOT / "screenshots"

# Annexure-1 Title Page provided by user (kept outside repo root)
ANNEXURE_TITLE_DOCX = (ROOT.parent / "SDP_TITLE_PAGE_for_BCA_UG_Annexure_1.docx").resolve()


CONFIG = {
    "project_title": "SignLens — Real-Time ASL Fingerspelling Recognition",
    "university": "Indus University",
    "department": "Department of Computer Science",
    "guides": [
        "Nirali Ma’am",
        "Bhargavi Ma’am",
    ],
    "students": [
        "Sundari — [Roll No.]",
        "Alnoor — [Roll No.]",
    ],
    "degree_line": "[Degree Name] in Computer Science",
    "academic_year": "2025–2026",
    "place": "Ahmedabad",
    "doc_date": date.today().strftime("%d/%m/%Y"),
}


FIGURES: list[tuple[str, Path]] = [
    ("Figure 4.1 Architecture Diagram", SCREENSHOTS_DIR / "architecture.png"),
    ("Figure 5.1 Dashboard Page", SCREENSHOTS_DIR / "dashboard.png"),
    ("Figure 5.2 Live Recognition Page", SCREENSHOTS_DIR / "live_recognition.png"),
    ("Figure 5.3 Word Suggestions", SCREENSHOTS_DIR / "suggestions.png"),
    ("Figure 5.4 Logs Page", SCREENSHOTS_DIR / "logs.png"),
    ("Figure 5.5 Settings Page", SCREENSHOTS_DIR / "settings.png"),
]


def _set_default_font(doc: Document, font_name: str = "Times New Roman", size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # Make headings clean (no theme colors). Word may still apply theme,
    # but enforcing font/size helps keep it professional.
    for heading_name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        try:
            st = doc.styles[heading_name]
        except KeyError:
            continue
        st.font.name = font_name
        st.font.size = Pt(size)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _center_bold(doc: Document, text: str, size_pt: int = 16) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size_pt)


def _para(doc: Document, text: str = "") -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _figure(doc: Document, caption: str, path: Path) -> None:
    _para(doc, caption)
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.0))
        cap = doc.add_paragraph(f"Caption: {caption}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[SCREENSHOT SLOT: {path.name} not found in {SCREENSHOTS_DIR}]" )
        if p.runs:
            p.runs[0].italic = True


def build_document() -> Document:
    doc = Document()
    _set_default_font(doc)

    # TITLE PAGE
    _center_bold(doc, "PROJECT REPORT", 18)
    _center_bold(doc, CONFIG["project_title"], 16)
    _para(doc)
    _para(doc, "Submitted in partial fulfillment of the requirements for the award of the degree of")
    _para(doc, CONFIG["degree_line"])
    _para(doc)
    _para(doc, "Submitted by:")
    for s in CONFIG["students"]:
        _para(doc, f"• {s}")
    _para(doc)
    _para(doc, "Under the guidance of:")
    for g in CONFIG["guides"]:
        _para(doc, f"• {g}")
    _para(doc)
    _para(doc, CONFIG["department"])
    _para(doc, CONFIG["university"])
    _para(doc, f"Academic Year: {CONFIG['academic_year']}")
    doc.add_page_break()

    # CERTIFICATE
    doc.add_heading("CERTIFICATE", level=1)
    student_names = ", ".join([s.split("—")[0].strip() for s in CONFIG["students"]])
    guide_names = ", ".join(CONFIG["guides"])
    _para(
        doc,
        f"This is to certify that the project entitled “{CONFIG['project_title']}” is a bonafide work carried out by "
        f"{student_names}, students of {CONFIG['department']}, {CONFIG['university']}, in partial fulfillment for the award "
        f"of the degree of {CONFIG['degree_line']} during the academic year {CONFIG['academic_year']}, under the guidance of {guide_names}."
    )
    _para(doc)
    _para(doc, "Guides: ____________________")
    _para(doc, "Head of Department: ____________________")
    _para(doc, f"Place: {CONFIG['place']}")
    _para(doc, f"Date: {CONFIG['doc_date']}")
    doc.add_page_break()

    # DECLARATION
    doc.add_heading("DECLARATION", level=1)
    _para(
        doc,
        f"We hereby declare that the project report entitled “{CONFIG['project_title']}” is an original work carried out by us under the guidance of "
        f"{guide_names}, and that it has not been submitted previously for the award of any degree or diploma in any institution or university."
    )
    _para(doc)
    _para(doc, "Signatures:")
    for s in CONFIG["students"]:
        _para(doc, f"{s.split('—')[0].strip()}: ____________________")
    _para(doc, f"Place: {CONFIG['place']}")
    _para(doc, f"Date: {CONFIG['doc_date']}")
    doc.add_page_break()

    # ACKNOWLEDGEMENT
    doc.add_heading("ACKNOWLEDGEMENT", level=1)
    _para(
        doc,
        "We express our sincere gratitude to our guides and the Department for continuous guidance and encouragement throughout this project. "
        "We also thank the Head of Department and faculty members for their support, and our friends and family for motivation and assistance during "
        "the completion of this work."
    )

    # ABSTRACT
    doc.add_heading("ABSTRACT", level=1)
    _para(
        doc,
        "SignLens is a real-time web-based American Sign Language (ASL) fingerspelling recognition system designed to recognize static ASL letters from live webcam video. "
        "The system uses MediaPipe HandLandmarker to detect 21 hand landmarks, transforms these landmarks into normalized features, and classifies them into 24 ASL letter classes "
        "(A–Y excluding J and Z, which require motion). To improve reliability, the system applies geometric disambiguation rules for visually similar hand shapes and uses temporal "
        "smoothing across recent predictions to reduce flicker. A word-building interface integrates word suggestions and text-to-speech to support end-to-end communication. "
        "The solution is implemented using a FastAPI + Socket.IO backend for real-time inference and a React (Vite) frontend for an interactive dashboard experience."
    )

    # TOC placeholder
    doc.add_heading("TABLE OF CONTENTS", level=1)
    _para(doc, "Note: In MS Word, generate the Table of Contents via References → Table of Contents (using Heading styles).")
    doc.add_page_break()

    # CHAPTER 1
    doc.add_heading("CHAPTER 1 — INTRODUCTION", level=1)
    doc.add_heading("1.1 Background", level=2)
    _para(
        doc,
        "American Sign Language enables communication for the Deaf and Hard-of-Hearing community. Fingerspelling is used to spell names, places, and words without dedicated gestures. "
        "Automating fingerspelling recognition using computer vision can support accessibility and assistive tools."
    )
    doc.add_heading("1.2 Problem Statement", level=2)
    _para(doc, "Design and implement a real-time system that recognizes ASL fingerspelled letters from a webcam feed, stabilizes predictions over time, and helps users form words and sentences.")
    doc.add_heading("1.3 Objectives", level=2)
    _bullets(
        doc,
        [
            "Detect a single hand in each frame and extract landmarks reliably.",
            "Classify static ASL letters into 24 classes (excluding dynamic gestures J and Z).",
            "Reduce confusion between similar signs using geometric disambiguation.",
            "Improve stability using temporal smoothing over multiple frames.",
            "Provide a web UI for live recognition, history, word suggestions, and logs.",
            "Provide REST + WebSocket interfaces for integration and observability.",
        ],
    )
    doc.add_heading("1.4 Scope", level=2)
    _para(doc, "Included: real-time static letter recognition, word building with suggestions, logs, and model status. Not included: dynamic gesture recognition for J and Z, multi-hand recognition, and full sign-language vocabulary beyond fingerspelling.")

    # CHAPTER 2
    doc.add_heading("CHAPTER 2 — LITERATURE SURVEY", level=1)
    doc.add_heading("2.1 Landmark-based Hand Gesture Recognition", level=2)
    _para(doc, "Landmark-based approaches represent the hand pose using keypoints, which reduces sensitivity to background and lighting and supports real-time inference with lightweight models.")
    doc.add_heading("2.2 CNN-based Image Classification", level=2)
    _para(doc, "CNN-based classification can be effective but typically requires more data and compute. It is often used as a fallback or in combination with landmarks.")
    doc.add_heading("2.3 Temporal Smoothing", level=2)
    _para(doc, "Real-time predictions can fluctuate frame-to-frame. Sliding window voting improves stability by requiring consistency across recent frames.")
    doc.add_heading("2.4 Rule-based Disambiguation", level=2)
    _para(doc, "Some ASL letters have visually similar shapes in static form. Geometric rules based on landmark relationships can reduce confusions between such letters.")

    # CHAPTER 3
    doc.add_heading("CHAPTER 3 — SYSTEM REQUIREMENTS", level=1)
    doc.add_heading("3.1 Hardware Requirements", level=2)
    _bullets(doc, ["Processor: Intel i5 / Ryzen 5 or above (recommended)", "RAM: 8 GB minimum (16 GB recommended)", "Webcam: Built-in or external (720p+ recommended)", "Storage: 2–5 GB free space"]) 
    doc.add_heading("3.2 Software Requirements", level=2)
    _bullets(doc, ["OS: Windows 10/11", "Python: 3.10+ (environment tested locally)", "Node.js: 18+", "Browser: Chrome/Edge"]) 
    doc.add_heading("3.3 Tools and Libraries", level=2)
    _para(doc, "Backend: FastAPI, python-socketio, Uvicorn, PyTorch, MediaPipe, OpenCV, NumPy, pyttsx3")
    _para(doc, "Frontend: React, Vite, TailwindCSS, socket.io-client, axios")

    # CHAPTER 4
    doc.add_heading("CHAPTER 4 — SYSTEM DESIGN", level=1)
    doc.add_heading("4.1 Overall Architecture", level=2)
    _para(
        doc,
        "Webcam Frame (Browser) → Socket.IO → Backend decodes frame → MediaPipe HandLandmarker (21 landmarks) → Feature extraction → Model inference → Disambiguation → Temporal smoothing → Predicted letter + confidence → Frontend display."
    )
    _figure(doc, "Figure 4.1 Architecture Diagram", SCREENSHOTS_DIR / "architecture.png")
    doc.add_heading("4.2 Module Design", level=2)
    _bullets(
        doc,
        [
            "Backend: REST API endpoints, Socket.IO inference stream, model loader, temporal smoothing, disambiguation, word predictor, and log buffer.",
            "Frontend: Dashboard, Live Recognition, Dataset Overview, Logs, Settings.",
        ],
    )

    # CHAPTER 5
    doc.add_heading("CHAPTER 5 — IMPLEMENTATION", level=1)
    doc.add_heading("5.1 Backend", level=2)
    _para(
        doc,
        "The backend is implemented using FastAPI and python-socketio. It loads a trained model on startup and exposes REST endpoints for health, model status, dataset info, logs, word suggestions, and text-to-speech. "
        "Real-time predictions are produced via the Socket.IO 'predict_frame' event, which accepts base64-encoded frames and emits 'prediction' events containing the predicted letter, confidence, and landmarks."
    )
    doc.add_heading("5.2 Frontend", level=2)
    _para(
        doc,
        "The frontend is built with React and Vite. The Live Recognition page streams frames to the backend and displays predicted letters, confidence, and overlays landmarks. It uses hold-frames and cooldown logic to accept stable letters into the current word."
    )

    for cap, p in FIGURES[1:]:
        _figure(doc, cap, p)

    # CHAPTER 6
    doc.add_heading("CHAPTER 6 — TESTING", level=1)
    _para(doc, "Testing covers API behavior, disambiguation rules, word suggestion functionality, and preprocessing utilities.")
    doc.add_heading("6.1 Sample Test Cases", level=2)
    _bullets(
        doc,
        [
            "GET /health returns status ok and model loaded flag.",
            "GET /model-status returns model type, classes, device, and load time.",
            "POST /suggest-words returns up to 4 suggestions for the current incomplete word.",
            "Live streaming via Socket.IO produces prediction events without errors.",
        ],
    )

    # CHAPTER 7
    doc.add_heading("CHAPTER 7 — RESULTS AND DISCUSSION", level=1)
    _para(doc, "The system provides real-time letter predictions and stable output through temporal smoothing. Word suggestions improve usability for sentence construction. Logs and status endpoints support debugging and monitoring.")

    # CHAPTER 8
    doc.add_heading("CHAPTER 8 — CONCLUSION AND FUTURE WORK", level=1)
    _para(doc, "SignLens demonstrates a practical real-time ASL fingerspelling recognition pipeline with a web interface and real-time backend.")
    doc.add_heading("8.1 Future Enhancements", level=2)
    _bullets(
        doc,
        [
            "Add dynamic gesture recognition for J and Z.",
            "Improve robustness under occlusion and extreme angles.",
            "Add multi-hand support.",
            "Optional persistence for sessions and analytics.",
        ],
    )

    # REFERENCES
    doc.add_heading("REFERENCES", level=1)
    _para(doc, "1) FastAPI Documentation")
    _para(doc, "2) MediaPipe HandLandmarker Documentation")
    _para(doc, "3) PyTorch Documentation")
    _para(doc, "4) Socket.IO Documentation")
    _para(doc, "5) React + Vite Documentation")

    # APPENDIX
    doc.add_heading("APPENDIX A — INSTALLATION AND RUN PROCEDURE (WINDOWS)", level=1)
    _para(doc, "Backend:")
    _bullets(doc, ["Create venv: python -m venv .venv", "Activate: .venv\\Scripts\\activate", "Install: pip install -r backend\\requirements.txt", "Run: python -m uvicorn backend.server:app --host 127.0.0.1 --port 8000 --reload"]) 
    _para(doc, "Frontend:")
    _bullets(doc, ["cd frontend", "npm install", "npm run dev"]) 
    _para(doc, "URLs:")
    _bullets(doc, ["Frontend: http://localhost:5173", "Backend Docs: http://127.0.0.1:8000/docs"]) 

    return doc


def _iter_body_elements(doc: Document):
    # Yields underlying XML elements in document body (paragraphs, tables, etc.)
    return list(doc.element.body)


def _element_text(el) -> str:
    # Best-effort extraction of all text inside a paragraph element.
    # For non-paragraph elements, returns empty string.
    tag = getattr(el, "tag", "")
    if not tag.endswith("}p"):
        return ""
    texts = []
    for node in el.iter():
        if getattr(node, "tag", "").endswith("}t") and node.text:
            texts.append(node.text)
    return "".join(texts).strip()


def _append_body_from(src: Document, dst: Document, start_index: int = 0) -> None:
    body = _iter_body_elements(src)
    for el in body[start_index:]:
        dst.element.body.append(deepcopy(el))


def build_document_with_annexure() -> Document:
    """Create final report:
    - Prepend Annexure-1 title page (exact formatting from DOCX)
    - Append the report content (skipping the auto-generated title page)
    """
    report = build_document()

    if not ANNEXURE_TITLE_DOCX.exists():
        # Fall back to report as-is if annexure file is missing.
        return report

    annexure = Document(str(ANNEXURE_TITLE_DOCX))
    final_doc = Document()
    _set_default_font(final_doc)

    # 1) Copy annexure title page content exactly.
    _append_body_from(annexure, final_doc, start_index=0)
    final_doc.add_page_break()

    # 2) Append report starting from CERTIFICATE (skip original generated title page).
    report_body = _iter_body_elements(report)
    start_idx = 0
    for i, el in enumerate(report_body):
        if _element_text(el).upper() == "CERTIFICATE":
            start_idx = i
            break
    _append_body_from(report, final_doc, start_index=start_idx)

    return final_doc


def main() -> int:
    # Backup current doc if it exists (prevents accidental loss)
    if OUT_PATH.exists():
        try:
            OUT_PATH.replace(BACKUP_PATH)
        except Exception:
            # If backup fails, continue generating anyway.
            pass

    doc = build_document_with_annexure()
    try:
        doc.save(str(OUT_PATH))
    except PermissionError:
        print(f"ERROR: Cannot write {OUT_PATH}. Close the document if it is open and try again.")
        return 2

    print(f"Generated: {OUT_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
